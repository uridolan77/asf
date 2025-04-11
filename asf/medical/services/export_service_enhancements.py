"""
Enhanced Export Service for the Medical Research Synthesizer.

This module provides enhancements to the Export Service, including:
- Support for more export formats
- Better error handling for export errors
- Validation of input data
- Progress tracking for large exports
"""

import os
import json
import csv
import tempfile
import logging
import time
import hashlib
import asyncio
from typing import Dict, List, Optional, Any, Union, Callable
from datetime import datetime
from enum import Enum

from asf.medical.core.exceptions import (
    ExportError, ValidationError, FileError
)
from asf.medical.core.progress_tracker import ProgressTracker

# Set up logging
logger = logging.getLogger(__name__)

class ExportFormat(str, Enum):
    """
    Supported export formats.
    """
    JSON = "json"
    CSV = "csv"
    EXCEL = "excel"
    PDF = "pdf"
    XML = "xml"
    HTML = "html"
    MARKDOWN = "markdown"
    BIBTEX = "bibtex"
    RIS = "ris"
    DOCX = "docx"

class ExportProgressTracker(ProgressTracker):
    """
    Progress tracker for export operations.
    
    This class extends the base ProgressTracker to provide export-specific
    progress tracking functionality.
    """
    
    def __init__(self, export_id: str, total_steps: int = 100):
        """
        Initialize the export progress tracker.
        
        Args:
            export_id: Export ID
            total_steps: Total number of steps in the export
        """
        super().__init__(operation_id=export_id, total_steps=total_steps)
        self.export_id = export_id
        self.export_format = "unknown"
        self.start_time = time.time()
        self.file_path = None
        
    def set_export_format(self, export_format: str):
        """
        Set the export format.
        
        Args:
            export_format: Format of the export
        """
        self.export_format = export_format
        
    def set_file_path(self, file_path: str):
        """
        Set the export file path.
        
        Args:
            file_path: Path to the exported file
        """
        self.file_path = file_path
        
    def get_progress_details(self) -> Dict[str, Any]:
        """
        Get detailed progress information.
        
        Returns:
            Dictionary with progress details
        """
        details = super().get_progress_details()
        details.update({
            "export_id": self.export_id,
            "export_format": self.export_format,
            "elapsed_time": time.time() - self.start_time,
            "file_path": self.file_path
        })
        return details

def validate_export_input(func):
    """
    Decorator for validating export input data.
    
    This decorator validates input parameters for export methods.
    """
    async def wrapper(self, *args, **kwargs):
        # Extract common parameters
        data = kwargs.get('data', {})
        export_format = kwargs.get('export_format', None)
        
        # Validate data
        if not data:
            raise ValidationError("Data cannot be empty")
            
        # Validate export format
        if export_format is not None:
            try:
                export_format = ExportFormat(export_format.lower())
            except ValueError:
                valid_formats = ", ".join([f.value for f in ExportFormat])
                raise ValidationError(f"Invalid export format: {export_format}. Valid formats: {valid_formats}")
                
        # Call the original function
        return await func(self, *args, **kwargs)
    return wrapper

def track_export_progress(export_format: str, total_steps: int = 100):
    """
    Decorator for tracking export progress.
    
    This decorator adds progress tracking to export methods.
    
    Args:
        export_format: Format of the export
        total_steps: Total number of steps in the export
    """
    def decorator(func):
        async def wrapper(self, *args, **kwargs):
            # Generate a deterministic export ID based on the function and parameters
            func_name = func.__name__
            param_str = f"{func_name}:{args}:{kwargs}"
            export_id = hashlib.md5(param_str.encode()).hexdigest()
            
            # Create progress tracker
            tracker = ExportProgressTracker(export_id, total_steps)
            tracker.set_export_format(export_format)
            
            # Initialize progress
            tracker.update(0, "Starting export")
            
            # Add tracker to kwargs
            kwargs['progress_tracker'] = tracker
            
            try:
                # Call the original function
                result = await func(self, *args, **kwargs)
                
                # Set file path if result is a string (file path)
                if isinstance(result, str):
                    tracker.set_file_path(result)
                
                # Mark as complete
                tracker.complete("Export completed successfully")
                
                return result
            except Exception as e:
                # Mark as failed
                tracker.fail(f"Export failed: {str(e)}")
                
                # Re-raise the exception
                raise
        return wrapper
    return decorator

def enhanced_export_error_handling(func):
    """
    Decorator for enhanced error handling in export methods.
    
    This decorator adds detailed error handling to export methods.
    """
    async def wrapper(self, *args, **kwargs):
        try:
            # Call the original function
            return await func(self, *args, **kwargs)
        except ValidationError:
            # Re-raise validation errors
            raise
        except FileError:
            # Re-raise file errors
            raise
        except ExportError:
            # Re-raise export errors
            raise
        except Exception as e:
            # Log unexpected errors
            logger.error(f"Unexpected error in {func.__name__}: {str(e)}", exc_info=True)
            
            # Convert to ExportError
            raise ExportError(
                format=kwargs.get('export_format', 'unknown'),
                message=f"Unexpected error: {str(e)}"
            )
    return wrapper

# Example usage:
"""
class EnhancedExportService:
    @validate_export_input
    @track_export_progress("json", total_steps=3)
    @enhanced_export_error_handling
    async def export_to_json(
        self,
        data: Dict[str, Any],
        include_abstracts: bool = True,
        include_metadata: bool = True,
        progress_tracker: Optional[ExportProgressTracker] = None
    ) -> str:
        # Implementation with progress tracking
        if progress_tracker:
            progress_tracker.update(1, "Filtering data")
            
        # Filter data if needed
        filtered_data = self._filter_data(data, include_abstracts, include_metadata)
        
        if progress_tracker:
            progress_tracker.update(2, "Creating JSON file")
            
        # Create a temporary file
        fd, temp_path = tempfile.mkstemp(suffix=".json")
        os.close(fd)
        
        # Write data to the file
        with open(temp_path, 'w') as f:
            json.dump(filtered_data, f, indent=2)
            
        if progress_tracker:
            progress_tracker.update(3, "Finalizing export")
            
        # Create a permanent file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"export_{timestamp}.json"
        file_path = os.path.join(self.export_dir, file_name)
        
        # Move the temporary file to the permanent location
        os.rename(temp_path, file_path)
        
        logger.info(f"Data exported to JSON: {file_path}")
        
        return file_path
"""

# Additional export format implementations:

def export_to_xml(data: Dict[str, Any], output_path: str) -> str:
    """
    Export data to XML.
    
    Args:
        data: Data to export
        output_path: Path to save the XML file
        
    Returns:
        Path to the exported file
    """
    try:
        import xml.etree.ElementTree as ET
        from xml.dom import minidom
    except ImportError:
        raise ImportError("xml.etree.ElementTree is required for XML export")
        
    # Create root element
    root = ET.Element("export")
    
    # Add metadata
    metadata = ET.SubElement(root, "metadata")
    if "query" in data:
        query_elem = ET.SubElement(metadata, "query")
        query_elem.text = data["query"]
    if "timestamp" in data:
        timestamp_elem = ET.SubElement(metadata, "timestamp")
        timestamp_elem.text = data["timestamp"]
        
    # Add results
    results_elem = ET.SubElement(root, "results")
    for result in data.get("results", []):
        result_elem = ET.SubElement(results_elem, "result")
        
        # Add result fields
        for key, value in result.items():
            # Skip complex fields
            if isinstance(value, (dict, list)):
                continue
                
            field_elem = ET.SubElement(result_elem, key)
            field_elem.text = str(value)
            
    # Convert to string and pretty-print
    xml_str = ET.tostring(root, encoding="utf-8")
    pretty_xml = minidom.parseString(xml_str).toprettyxml(indent="  ")
    
    # Write to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(pretty_xml)
        
    return output_path

def export_to_html(data: Dict[str, Any], output_path: str, template: Optional[str] = None) -> str:
    """
    Export data to HTML.
    
    Args:
        data: Data to export
        output_path: Path to save the HTML file
        template: Path to HTML template file (optional)
        
    Returns:
        Path to the exported file
    """
    try:
        from jinja2 import Template
    except ImportError:
        raise ImportError("jinja2 is required for HTML export")
        
    # Default template
    default_template = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Medical Research Synthesizer Export</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 20px; }
            h1 { color: #2c3e50; }
            .result { margin-bottom: 20px; padding: 10px; border: 1px solid #ddd; }
            .title { font-weight: bold; color: #3498db; }
            .abstract { color: #555; }
            .metadata { color: #7f8c8d; font-size: 0.9em; }
        </style>
    </head>
    <body>
        <h1>Medical Research Synthesizer Export</h1>
        {% if query %}
        <p><strong>Query:</strong> {{ query }}</p>
        {% endif %}
        <p><strong>Results:</strong> {{ results|length }}</p>
        
        <div class="results">
        {% for result in results %}
            <div class="result">
                <div class="title">{{ result.title }}</div>
                <div class="metadata">
                    {% if result.authors %}
                    <p><strong>Authors:</strong> {{ result.authors|join(', ') }}</p>
                    {% endif %}
                    {% if result.journal %}
                    <p><strong>Journal:</strong> {{ result.journal }}</p>
                    {% endif %}
                    {% if result.publication_date %}
                    <p><strong>Date:</strong> {{ result.publication_date }}</p>
                    {% endif %}
                </div>
                {% if result.abstract %}
                <div class="abstract">
                    <p><strong>Abstract:</strong> {{ result.abstract }}</p>
                </div>
                {% endif %}
            </div>
        {% endfor %}
        </div>
    </body>
    </html>
    """
    
    # Load template
    if template and os.path.exists(template):
        with open(template, "r") as f:
            template_str = f.read()
    else:
        template_str = default_template
        
    # Create template
    template = Template(template_str)
    
    # Render template
    html = template.render(
        query=data.get("query", ""),
        results=data.get("results", []),
        timestamp=data.get("timestamp", datetime.now().isoformat())
    )
    
    # Write to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
        
    return output_path

def export_to_markdown(data: Dict[str, Any], output_path: str) -> str:
    """
    Export data to Markdown.
    
    Args:
        data: Data to export
        output_path: Path to save the Markdown file
        
    Returns:
        Path to the exported file
    """
    # Create markdown content
    lines = []
    
    # Add title
    lines.append("# Medical Research Synthesizer Export\n")
    
    # Add metadata
    if "query" in data:
        lines.append(f"**Query:** {data['query']}\n")
    if "timestamp" in data:
        lines.append(f"**Date:** {data['timestamp']}\n")
        
    # Add results
    results = data.get("results", [])
    lines.append(f"**Results:** {len(results)}\n")
    
    for i, result in enumerate(results, 1):
        lines.append(f"## {i}. {result.get('title', 'Untitled')}\n")
        
        # Add authors
        authors = result.get("authors", [])
        if authors:
            if isinstance(authors, list):
                authors_str = ", ".join(authors)
            else:
                authors_str = authors
            lines.append(f"**Authors:** {authors_str}\n")
            
        # Add journal and date
        journal = result.get("journal", "")
        if journal:
            lines.append(f"**Journal:** {journal}\n")
            
        date = result.get("publication_date", "")
        if date:
            lines.append(f"**Date:** {date}\n")
            
        # Add abstract
        abstract = result.get("abstract", "")
        if abstract:
            lines.append(f"**Abstract:**\n\n{abstract}\n")
            
        lines.append("\n---\n")
        
    # Write to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
        
    return output_path

def export_to_bibtex(data: Dict[str, Any], output_path: str) -> str:
    """
    Export data to BibTeX.
    
    Args:
        data: Data to export
        output_path: Path to save the BibTeX file
        
    Returns:
        Path to the exported file
    """
    # Create BibTeX content
    lines = []
    
    # Add comment with metadata
    lines.append("% Medical Research Synthesizer Export")
    if "query" in data:
        lines.append(f"% Query: {data['query']}")
    if "timestamp" in data:
        lines.append(f"% Date: {data['timestamp']}")
    lines.append("")
    
    # Add entries
    for result in data.get("results", []):
        # Create BibTeX key
        pmid = result.get("pmid", "")
        first_author = ""
        authors = result.get("authors", [])
        if authors and isinstance(authors, list) and len(authors) > 0:
            first_author = authors[0].split()[-1] if " " in authors[0] else authors[0]
        year = result.get("publication_date", "")[:4] if result.get("publication_date", "") else ""
        
        if pmid and first_author and year:
            key = f"{first_author}{year}{pmid}"
        elif pmid:
            key = f"pmid{pmid}"
        else:
            key = f"entry{hash(result.get('title', ''))}"
            
        # Start entry
        lines.append(f"@article{{{key},")
        
        # Add fields
        if "title" in result:
            lines.append(f"  title = {{{result['title']}}},")
            
        if authors:
            if isinstance(authors, list):
                authors_str = " and ".join(authors)
            else:
                authors_str = authors
            lines.append(f"  author = {{{authors_str}}},")
            
        if "journal" in result:
            lines.append(f"  journal = {{{result['journal']}}},")
            
        if "publication_date" in result:
            year = result["publication_date"][:4] if result["publication_date"] else ""
            if year:
                lines.append(f"  year = {{{year}}},")
                
        if "volume" in result:
            lines.append(f"  volume = {{{result['volume']}}},")
            
        if "issue" in result:
            lines.append(f"  number = {{{result['issue']}}},")
            
        if "pages" in result:
            lines.append(f"  pages = {{{result['pages']}}},")
            
        if "doi" in result:
            lines.append(f"  doi = {{{result['doi']}}},")
            
        if "pmid" in result:
            lines.append(f"  pmid = {{{result['pmid']}}},")
            
        # End entry (remove trailing comma from last field)
        lines[-1] = lines[-1].rstrip(",")
        lines.append("}")
        lines.append("")
        
    # Write to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
        
    return output_path

def export_to_ris(data: Dict[str, Any], output_path: str) -> str:
    """
    Export data to RIS (Research Information Systems) format.
    
    Args:
        data: Data to export
        output_path: Path to save the RIS file
        
    Returns:
        Path to the exported file
    """
    # Create RIS content
    lines = []
    
    # Add entries
    for result in data.get("results", []):
        # Start entry
        lines.append("TY  - JOUR")  # Type: Journal Article
        
        # Add fields
        if "title" in result:
            lines.append(f"TI  - {result['title']}")
            
        # Add authors
        authors = result.get("authors", [])
        if isinstance(authors, list):
            for author in authors:
                lines.append(f"AU  - {author}")
        elif authors:
            lines.append(f"AU  - {authors}")
            
        if "journal" in result:
            lines.append(f"JO  - {result['journal']}")
            
        if "publication_date" in result:
            date = result["publication_date"]
            year = date[:4] if date else ""
            if year:
                lines.append(f"PY  - {year}")
                
        if "volume" in result:
            lines.append(f"VL  - {result['volume']}")
            
        if "issue" in result:
            lines.append(f"IS  - {result['issue']}")
            
        if "pages" in result:
            lines.append(f"SP  - {result['pages']}")
            
        if "doi" in result:
            lines.append(f"DO  - {result['doi']}")
            
        if "pmid" in result:
            lines.append(f"AN  - {result['pmid']}")
            
        if "abstract" in result:
            lines.append(f"AB  - {result['abstract']}")
            
        # End entry
        lines.append("ER  - ")
        lines.append("")
        
    # Write to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
        
    return output_path

def export_to_docx(data: Dict[str, Any], output_path: str) -> str:
    """
    Export data to DOCX.
    
    Args:
        data: Data to export
        output_path: Path to save the DOCX file
        
    Returns:
        Path to the exported file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise ImportError("python-docx is required for DOCX export")
        
    # Create document
    doc = Document()
    
    # Add title
    doc.add_heading("Medical Research Synthesizer Export", level=1)
    
    # Add metadata
    if "query" in data:
        doc.add_paragraph(f"Query: {data['query']}")
    if "timestamp" in data:
        doc.add_paragraph(f"Date: {data['timestamp']}")
        
    # Add results
    results = data.get("results", [])
    doc.add_paragraph(f"Results: {len(results)}")
    
    for i, result in enumerate(results, 1):
        # Add title
        doc.add_heading(f"{i}. {result.get('title', 'Untitled')}", level=2)
        
        # Add authors
        authors = result.get("authors", [])
        if authors:
            if isinstance(authors, list):
                authors_str = ", ".join(authors)
            else:
                authors_str = authors
            doc.add_paragraph(f"Authors: {authors_str}")
            
        # Add journal and date
        journal = result.get("journal", "")
        if journal:
            doc.add_paragraph(f"Journal: {journal}")
            
        date = result.get("publication_date", "")
        if date:
            doc.add_paragraph(f"Date: {date}")
            
        # Add abstract
        abstract = result.get("abstract", "")
        if abstract:
            doc.add_paragraph("Abstract:")
            doc.add_paragraph(abstract)
            
        # Add separator
        doc.add_paragraph("---")
        
    # Save document
    doc.save(output_path)
    
    return output_path
